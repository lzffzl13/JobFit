from io import BytesIO

from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def _extension(filename: str) -> str:
    if "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[-1].lower()


async def parse_upload(file: UploadFile) -> str:
    content = await file.read()
    ext = _extension(file.filename or "")

    if ext == ".pdf":
        return parse_pdf(content)
    if ext == ".docx":
        return parse_docx(content)
    if ext in {".txt", ".md"}:
        return parse_text(content)

    supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
    raise ValueError(f"Unsupported file type '{ext or 'unknown'}'. Supported: {supported}")


def parse_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            pages.append(f"[Page {index}]\n{text}")
    return "\n\n".join(pages).strip()


def parse_docx(content: bytes) -> str:
    document = Document(BytesIO(content))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    table_lines: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))

    return "\n".join([*paragraphs, *table_lines]).strip()


def parse_text(content: bytes) -> str:
    for encoding in ("utf-8", "gb18030"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore").strip()
